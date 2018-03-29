import requests, json
from docx import Document
from docx.shared import RGBColor

def API_Response(API_key, Transcript_ID):
	api_url='https://api.capio.ai/v1/speech/transcript/'+ Transcript_ID
	headers={'apikey':API_key}
	Response=requests.get(api_url,headers=headers)
	return Response

def write_transcript_to_file(Resp):
        transcript_doc = Document()
        transcript_doc.add_heading('TRANSCRIPT')
        for sentence in Resp:
            
            line = transcript_doc.add_paragraph() 
            t=sentence['result'][0]['alternative'][0]['words'][0]['from']
            time_stamp = '{0:02.0f}:{1:02.0f}'.format(*divmod(t * 60, 60))
            line.add_run(' ['+ time_stamp +'] : ')
            for word in sentence['result'][0]['alternative'][0]['words']:
                w_add = line.add_run(word['word'] + ' ')
                if word['confidence'] < 0.75:
                    font=w_add.font
                    font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        transcript_doc.save('transcript.docx')
        return True


resp=API_Response('262ac9a0c9ba4d179aad4c0b9b02120a', '593f237fbcae700012ba8fcd')
print(resp)
if resp.status_code==200:
    resp=resp.json()
    write_transcript_to_file(resp)
